# -*- coding: utf-8 -*-
"""Exportación de informes MDeIA a Excel y Word."""

from __future__ import annotations

from datetime import date
from io import BytesIO
from typing import Any

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from constants import FASE1_TITULO
from lib.informe import generar_informe_html
from lib.mdeia_model import (
    areas_df,
    brechas_prioritarias,
    calcular_imd,
    detalle_evaluacion,
    guia_indices,
    load_framework,
    pilot_codigos,
)


def _meta_informe_resuelto(meta: dict | None) -> dict:
    """Unifica unidad académica en meta (label + sede) para informes."""
    from lib.unidades import unidad_label

    meta = dict(meta or {})
    meta.setdefault("fecha", date.today().isoformat())
    meta.setdefault("responsable", "Observatorio de IA · UCCuyo")
    uid = meta.get("unidad_id")
    if uid and not meta.get("unidad_label"):
        meta["unidad_label"] = unidad_label(uid)
    label = meta.get("unidad_label") or meta.get("sede") or "Institucional"
    meta["unidad_label"] = label
    meta["sede"] = label
    return meta


def _meta_default(meta: dict | None) -> dict:
    return _meta_informe_resuelto(meta)


def _agregar_guia_indices_word(doc: Document, resultado: dict) -> None:
    doc.add_paragraph(
        "Los índices se calculan sobre los indicadores respondidos en esta carga. "
        "Rango habitual: 0 % (ninguno satisface) a 100 % (todos satisface)."
    )
    for item in guia_indices(resultado):
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{item['nombre']}: {item['valor']} ")
        run.bold = True
        p.add_run(f"(rango {item['rango']}). {item['interpretacion']}")


def _agregar_brechas_word(doc: Document, brechas: pd.DataFrame) -> None:
    doc.add_paragraph(
        "Indicadores evaluados que aún no alcanzan el umbral de madurez. "
        "Prioridad mayor = brecha más grande (p. ej. nivel 0 vs meta 3)."
    )
    if brechas.empty:
        doc.add_paragraph("No se registran brechas en los indicadores evaluados.")
        return
    for _, row in brechas.iterrows():
        p = doc.add_paragraph(style="List Number")
        tit = p.add_run(str(row["Indicador"]))
        tit.bold = True
        p.add_run(
            f"\nCódigo: {row['Código']} · Área: {row['Área']} · Reto: {row['Reto']}"
        )
        p.add_run(
            f"\nValor actual: {row['Valor actual']} · Meta: {row['Meta']}"
        )
        interp = p.add_run(f"\n{row['Interpretación']}")
        interp.italic = True


def generar_excel_bytes(
    respuestas: dict[str, Any],
    *,
    meta: dict | None = None,
    modo: str = "piloto",
) -> bytes:
    codigos = pilot_codigos() if modo == "piloto" else None
    resultado = calcular_imd(respuestas, codigos=codigos)
    meta = _meta_default(meta)
    fw = load_framework()
    area_names = {a["id"]: a["nombre"] for a in areas_df().to_dict("records")}

    resumen = pd.DataFrame(
        [
            {"Indicador": "Unidad académica", "Valor": meta["unidad_label"]},
            {"Indicador": "Fecha", "Valor": meta.get("fecha", "")},
            {"Indicador": "Responsable", "Valor": meta.get("responsable", "")},
        ]
    )
    guia = pd.DataFrame(guia_indices(resultado))
    guia = guia.rename(
        columns={
            "nombre": "Índice",
            "valor": "Valor",
            "rango": "Rango",
            "interpretacion": "Interpretación",
        }
    )

    df_area = pd.DataFrame(resultado["por_area"])
    if not df_area.empty:
        df_area["Área"] = df_area["grupo"].map(area_names).fillna(df_area["grupo"])
        df_area = df_area.rename(
            columns={"ratio_pct": "IMD (%)", "satisfechos": "Satisfechos", "evaluados": "Evaluados"}
        )[["Área", "IMD (%)", "Satisfechos", "Evaluados"]]

    df_reto = pd.DataFrame(resultado["por_reto"])
    if not df_reto.empty:
        df_reto = df_reto.rename(
            columns={"grupo": "Reto", "ratio_pct": "IMD (%)", "satisfechos": "Satisfechos", "evaluados": "Evaluados"}
        )

    det = detalle_evaluacion(respuestas)
    if codigos is not None and not det.empty:
        det = det[det["Código"].isin(codigos)]

    brechas = brechas_prioritarias(respuestas, top_n=15)
    if codigos is not None and not brechas.empty:
        brechas = brechas[brechas["Código"].isin(codigos)]

    buf = BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        resumen.to_excel(writer, index=False, sheet_name="Resumen")
        guia.to_excel(writer, index=False, sheet_name="Índices")
        if not df_area.empty:
            df_area.to_excel(writer, index=False, sheet_name="Por área")
        if not df_reto.empty:
            df_reto.to_excel(writer, index=False, sheet_name="Por reto")
        if not det.empty:
            det.to_excel(writer, index=False, sheet_name="Detalle")
        if not brechas.empty:
            brechas.to_excel(writer, index=False, sheet_name="Brechas")
        pd.DataFrame([{"modelo": fw["meta"]["nombre"], "formula": resultado["formula"]}]).to_excel(
            writer, index=False, sheet_name="Metadatos"
        )
    buf.seek(0)
    return buf.getvalue()


def generar_word_bytes(
    respuestas: dict[str, Any],
    *,
    meta: dict | None = None,
    modo: str = "piloto",
) -> bytes:
    codigos = pilot_codigos() if modo == "piloto" else None
    resultado = calcular_imd(respuestas, codigos=codigos)
    meta = _meta_default(meta)
    fw = load_framework()
    titulo_modo = FASE1_TITULO if modo == "piloto" else "Catálogo completo"
    area_names = {a["id"]: a["nombre"] for a in areas_df().to_dict("records")}

    doc = Document()
    verde = RGBColor(4, 74, 48)

    titulo = doc.add_heading("Informe MDeIA UCCuyo", level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    run = sub.add_run(f"{titulo_modo} · {fw['meta']['institucion']}")
    run.bold = True
    run.font.color.rgb = verde

    doc.add_paragraph(
        f"Fecha: {meta.get('fecha')} · Responsable: {meta.get('responsable')}\n"
        f"Unidad académica: {meta['unidad_label']}"
    )

    doc.add_heading("Índices principales e interpretación", level=1)
    _agregar_guia_indices_word(doc, resultado)

    umbral = resultado["umbral_regional_objetivo"]
    diag = doc.add_paragraph()
    diag.add_run(
        f"Diagnóstico: {'Por encima' if resultado['imd'] >= umbral else 'Por debajo'} "
        f"del umbral regional ({umbral} %)."
    ).italic = True

    doc.add_heading("IMD por área", level=1)
    df_area = pd.DataFrame(resultado["por_area"])
    if not df_area.empty:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, text in enumerate(["Área", "IMD (%)", "Satisfechos", "Evaluados"]):
            hdr[i].text = text
        for _, row in df_area.iterrows():
            cells = table.add_row().cells
            area = area_names.get(row["grupo"], row["grupo"])
            cells[0].text = str(area)
            cells[1].text = str(row["ratio_pct"])
            cells[2].text = str(row["satisfechos"])
            cells[3].text = str(row["evaluados"])
    else:
        doc.add_paragraph("Sin datos por área.")

    doc.add_heading("Brechas prioritarias", level=1)
    brechas = brechas_prioritarias(respuestas, top_n=10)
    if codigos is not None and not brechas.empty:
        brechas = brechas[brechas["Código"].isin(codigos)]
    _agregar_brechas_word(doc, brechas)

    doc.add_heading("Detalle de indicadores", level=1)
    det = detalle_evaluacion(respuestas)
    if codigos is not None and not det.empty:
        det = det[det["Código"].isin(codigos)]
    if det.empty:
        doc.add_paragraph("Sin indicadores evaluados.")
    else:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        for i, text in enumerate(["Código", "Indicador", "Valor", "Satisface"]):
            table.rows[0].cells[i].text = text
        for _, row in det.head(40).iterrows():
            cells = table.add_row().cells
            cells[0].text = str(row["Código"])
            cells[1].text = str(row["Indicador"])
            cells[2].text = str(row["Valor"])
            cells[3].text = str(row["Satisface"])

    doc.add_paragraph()
    pie = doc.add_paragraph(f"Generado por {fw['meta']['nombre']} · {date.today().isoformat()}")
    pie.runs[0].font.size = Pt(9)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def generar_html_bytes(
    respuestas: dict[str, Any],
    *,
    meta: dict | None = None,
    modo: str = "piloto",
) -> str:
    return generar_informe_html(respuestas, modo=modo, meta_encuesta=_meta_default(meta))
